import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document

def extract_text_from_url(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            # Extract all paragraphs from the webpage
            paragraphs = soup.find_all('p')
            # Combine text from paragraphs up to 500 words
            word_count = 0
            extracted_text = []
            for p in paragraphs:
                if word_count >= 500:
                    break
                text = p.get_text().strip()
                words = text.split()
                word_count += len(words)
                if word_count <= 500:
                    extracted_text.append(text)
            return ' '.join(extracted_text)
        else:
            print(f"Failed to retrieve content from URL: {url}")
            return None
    except Exception as e:
        print(f"Error occurred while extracting text: {str(e)}")
        return None

def main(input_file):
    df = pd.read_excel(input_file, header=None)  # Assuming links are in the first column
    for i, row in df.iterrows():
        link = row[0]
        extracted_text = extract_text_from_url(link)
        if extracted_text:
            # Create a new document
            doc = Document()
            doc.add_heading(f"Document {i+1}", 0)
            doc.add_paragraph(extracted_text)
            doc_file_name = f"Document {i+1}.docx"
            doc.save(doc_file_name)
            print(f"Text extracted from {link} and saved to {doc_file_name}")

if __name__ == "__main__":
    input_excel_file = "lifestyleAndHobbies.xlsx"  # Provide the path to your input Excel file
    main(input_excel_file)
