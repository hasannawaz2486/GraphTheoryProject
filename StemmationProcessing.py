import os
import docx
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
import string

# Ensure NLTK stopwords are downloaded
import nltk


# Initialize Porter Stemmer
stemmer = PorterStemmer()

# Function to remove stopwords, punctuation, and perform stemming
def process_text(text):
    stop_words = set(stopwords.words('english'))
    word_tokens = word_tokenize(text)
    filtered_text = [word for word in word_tokens if word.lower() not in stop_words and word not in string.punctuation]
    stemmed_text = [stemmer.stem(word) for word in filtered_text]
    return ' '.join(stemmed_text)

# Function to process .docx files in a directory and save processed files into a new folder
def process_files(input_folder, output_folder):
    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Process each .docx file in the input folder
    for filename in os.listdir(input_folder):
        if filename.endswith('.docx'):  
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename)

            # Open the .docx file
            doc = docx.Document(input_path)

            # Extract text from the document
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Process the text (remove stopwords, punctuation, and perform stemming)
            processed_text = process_text(text)

            # Create a new document
            new_doc = docx.Document()
            new_doc.add_paragraph(processed_text)

            # Save the new document
            new_doc.save(output_path)

# Specify input and output folders
input_folder = 'ProcessScience& Education'  # Relative path from the current working directory
output_folder = 'StemmationProcessScience& Education'  # Relative path from the current working directory

# Get the absolute paths
current_directory = os.path.dirname(os.path.abspath(__file__))
input_folder = os.path.join(current_directory, input_folder)
output_folder = os.path.join(current_directory, output_folder)

# Process files
process_files(input_folder, output_folder)
